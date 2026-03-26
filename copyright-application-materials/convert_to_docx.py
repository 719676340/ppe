#!/usr/bin/env python3
"""
将软著申请材料Markdown文件转换为Word文档
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='宋体', font_size=12, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_title(doc, text, level=1):
    """添加标题"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_chinese_font(run, '黑体', 16 if level == 1 else 14, bold=True)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.space_before = Pt(12)
    paragraph.space_after = Pt(6)
    return paragraph

def add_normal_paragraph(doc, text, bold=False):
    """添加普通段落"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_chinese_font(run, '宋体', 12, bold=bold)
    paragraph.space_before = Pt(0)
    paragraph.space_after = Pt(6)
    return paragraph

def convert_application_form(doc, content):
    """转换软件著作权登记申请表"""
    # 标题
    title = doc.add_paragraph()
    run = title.add_run('软件著作权登记申请表')
    set_chinese_font(run, '黑体', 18, bold=True)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.space_after = Pt(12)

    lines = content.split('\n')
    in_table = False

    for line in lines:
        line = line.strip()

        if not line or line == '---':
            continue

        # 处理Markdown表格
        if line.startswith('|'):
            # 跳过表格分隔线
            if '|---' in line:
                in_table = True
                continue

            parts = [p.strip().replace('**', '') for p in line.split('|')[1:-1]]
            if len(parts) >= 2:
                label = parts[0]
                value = parts[1]

                # 跳过表头
                if label == '项目':
                    continue

                if value and value not in ['内容', '（待填写）']:
                    p = doc.add_paragraph()
                    run1 = p.add_run(f'{label}: ')
                    set_chinese_font(run1, '宋体', 12, bold=True)
                    run2 = p.add_run(value)
                    set_chinese_font(run2, '宋体', 12)
                    p.space_before = Pt(3)
                    p.space_after = Pt(3)

        # 处理标题
        elif line.startswith('## '):
            title_text = line.replace('## ', '').replace('**', '')
            add_title(doc, title_text, level=1)
        elif line.startswith('### '):
            title_text = line.replace('### ', '').replace('**', '')
            add_title(doc, title_text, level=2)
        # 处理列表
        elif line.startswith('- ') or line.startswith('**') or re.match(r'^\d+\.', line):
            add_normal_paragraph(doc, line.replace('**', ''))
        # 处理声明部分
        elif line.startswith('申请人声明'):
            add_normal_paragraph(doc, line, bold=True)
        elif line.startswith('申请人签字') or line.startswith('日期'):
            add_normal_paragraph(doc, line)

def convert_code_doc(doc, content):
    """转换源代码文档"""
    lines = content.split('\n')

    for line in lines:
        # 文件头
        if line.startswith('## ') and ('文件' in line or '=' in line):
            p = doc.add_paragraph()
            run = p.add_run(line.replace('## ', ''))
            set_chinese_font(run, '黑体', 12, bold=True)
            p.space_before = Pt(6)
            p.space_after = Pt(3)
        # 页码标记
        elif line.startswith('--- 第') and '页 ---' in line:
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_chinese_font(run, '宋体', 10)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 代码行
        elif line.strip():
            add_normal_paragraph(doc, line.rstrip())

def convert_general_doc(doc, content):
    """转换普通文档（用户手册、设计说明书）"""
    lines = content.split('\n')

    for line in lines:
        line = line.rstrip()

        if not line or line == '---':
            if not line:
                doc.add_paragraph()
            continue

        # 一级标题
        if line.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            set_chinese_font(run, '黑体', 18, bold=True)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.space_after = Pt(12)
        # 二级标题
        elif line.startswith('## '):
            add_title(doc, line[3:], level=1)
        # 三级标题
        elif line.startswith('### '):
            add_title(doc, line[4:], level=2)
        # 四级标题
        elif line.startswith('#### '):
            add_title(doc, line[5:], level=3)
        # 表格行 - 简化处理
        elif line.startswith('|') and '|---' not in line:
            parts = [p.strip().replace('**', '') for p in line.split('|')[1:-1]]
            if parts and any(parts):
                text = '  '.join([p for p in parts if p])
                add_normal_paragraph(doc, text)
        # 列表
        elif line.strip().startswith(('- ', '* ', '1. ', '2. ', '3. ', '4. ', '5. ',
                                      '6. ', '7. ', '8. ', '9. ')):
            add_normal_paragraph(doc, line.strip())
        # 普通段落
        elif line.strip():
            # 处理粗体标记
            text = line.replace('**', '').strip()
            add_normal_paragraph(doc, text)

def convert_markdown_to_docx(md_file, output_file, doc_type='general'):
    """将Markdown文件转换为Word文档"""
    doc = Document()

    # 设置页面边距（3cm约等于1.18英寸）
    section = doc.sections[0]
    section.top_margin = Inches(1.18)
    section.bottom_margin = Inches(1.18)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(1.18)

    # 读取Markdown内容
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    if doc_type == 'application_form':
        convert_application_form(doc, content)
    elif doc_type == 'code':
        convert_code_doc(doc, content)
    else:
        convert_general_doc(doc, content)

    # 保存文档
    doc.save(output_file)
    print(f'✓ 已生成: {output_file}')

def main():
    """主函数"""
    base_dir = Path('/Users/heqijie/个人项目/中烟软著/头盔/copyright-application-materials')

    print('开始转换Markdown文件到Word文档...\n')

    # 定义文件映射
    files = [
        ('软件著作权登记申请表.md', '软件著作权登记申请表.docx', 'application_form'),
        ('用户手册.md', '用户手册.docx', 'general'),
        ('设计说明书.md', '设计说明书.docx', 'general'),
        ('源代码文档.md', '源代码文档.docx', 'code'),
    ]

    success_count = 0
    for md_file, docx_file, doc_type in files:
        md_path = base_dir / md_file
        docx_path = base_dir / docx_file

        if md_path.exists():
            try:
                convert_markdown_to_docx(md_path, docx_path, doc_type)
                success_count += 1
            except Exception as e:
                print(f'✗ 转换失败 {md_file}: {e}')
        else:
            print(f'✗ 文件不存在: {md_path}')

    print(f'\n转换完成！成功生成 {success_count}/{len(files)} 个Word文档')
    print(f'输出目录: {base_dir}')

if __name__ == '__main__':
    main()
